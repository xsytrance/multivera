from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import chromadb
from chromadb.api.models.Collection import Collection
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
KNOWLEDGE_DIR = BASE_DIR / "knowledge" / "hackermouth"
CHROMA_DIR = BASE_DIR / ".chroma"
COLLECTION_NAME = "hackermouth_lore"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 100


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    step = max(1, chunk_size - overlap)

    while start < len(normalized):
        end = min(len(normalized), start + chunk_size)
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start += step

    return chunks


def get_chroma_client() -> chromadb.PersistentClient:
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    return chromadb.PersistentClient(path=str(CHROMA_DIR))


def get_collection() -> Collection:
    client = get_chroma_client()
    return client.get_or_create_collection(name=COLLECTION_NAME)


def rebuild_collection() -> Collection:
    client = get_chroma_client()
    try:
        client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass
    return client.get_or_create_collection(name=COLLECTION_NAME)


def ingest_docx_files(knowledge_dir: Path = KNOWLEDGE_DIR) -> dict[str, Any]:
    collection = rebuild_collection()
    docx_files = sorted(knowledge_dir.glob("*.docx"))

    documents: list[str] = []
    ids: list[str] = []
    metadatas: list[dict[str, Any]] = []
    files_processed = 0

    for path in docx_files:
        text = extract_docx_text(path)
        chunks = chunk_text(text)
        if not chunks:
            continue

        files_processed += 1
        for index, chunk in enumerate(chunks, start=1):
            documents.append(chunk)
            ids.append(f"{path.stem}::chunk_{index}")
            metadatas.append(
                {
                    "source_file": path.name,
                    "chunk_index": index,
                    "character_count": len(chunk),
                }
            )

    if documents:
        collection.upsert(documents=documents, ids=ids, metadatas=metadatas)

    return {
        "knowledge_dir": str(knowledge_dir),
        "files_found": len(docx_files),
        "files_processed": files_processed,
        "chunks_stored": len(documents),
        "collection_name": COLLECTION_NAME,
        "chroma_path": str(CHROMA_DIR),
    }


def search_lore(query: str, n_results: int = 3) -> list[dict[str, Any]]:
    if not query.strip():
        return []

    collection = get_collection()
    collection_count = collection.count()
    if collection_count == 0:
        return []

    result_count = min(max(1, n_results), collection_count)
    results = collection.query(query_texts=[query], n_results=result_count)

    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]
    distances = results.get("distances", [[]])[0]

    matches: list[dict[str, Any]] = []
    for index, document in enumerate(documents):
        metadata = metadatas[index] if index < len(metadatas) else {}
        distance = distances[index] if index < len(distances) else None
        matches.append(
            {
                "text": document,
                "metadata": metadata or {},
                "distance": distance,
            }
        )

    return matches
