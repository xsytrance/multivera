"""MultiVera RAG service — generalized ChromaDB lore ingestion and search.

Refactors hackermouth_rag.py to support:
- Per-project Chroma collections
- Multiple file formats (.txt, .md, .pdf, .docx)
- Source references and project_id in metadata
- Intelligent chunking with overlap
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import TYPE_CHECKING, Any, Dict, List, Optional

if TYPE_CHECKING:
    from chromadb.api.models.Collection import Collection

try:
    import chromadb
except ModuleNotFoundError:
    chromadb = None  # type: ignore[assignment]

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None  # type: ignore[misc,assignment]

try:
    import PyPDF2
except ModuleNotFoundError:
    PyPDF2 = None  # type: ignore[misc]

from backend.config import CHROMA_DIR, CHUNK_OVERLAP, CHUNK_SIZE

logger = logging.getLogger("multivera.rag_service")


def _ensure_chroma() -> None:
    if chromadb is None:
        raise RuntimeError("chromadb is not installed. Run: pip install chromadb")


def _collection_name_for_project(project_id: int) -> str:
    """Generate a Chroma collection name for a project."""
    return f"project_{project_id}_lore"


def normalize_text(text: str) -> str:
    """Collapse whitespace and strip."""
    return re.sub(r"\s+", " ", text).strip()


def extract_docx_text(path: Path) -> str:
    """Extract text from a .docx file."""
    if Document is None:
        raise RuntimeError("python-docx is not installed")
    document = Document(path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_pdf_text(path: Path) -> str:
    """Extract text from a .pdf file."""
    if PyPDF2 is None:
        raise RuntimeError("PyPDF2 is not installed")
    text_parts: List[str] = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(normalize_text(page_text))
    return "\n\n".join(text_parts)


def extract_txt_text(path: Path) -> str:
    """Extract text from a plain text or markdown file."""
    with open(path, "r", encoding="utf-8") as f:
        return normalize_text(f.read())


def extract_text_from_file(path: Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return extract_txt_text(path)
    raise ValueError(f"Unsupported file format: {suffix}")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Chunk text into overlapping segments."""
    normalized = normalize_text(text)
    if not normalized:
        return []

    chunks: List[str] = []
    start = 0
    step = max(1, chunk_size - overlap)

    while start < len(normalized):
        end = min(len(normalized), start + chunk_size)
        # Try to break at a sentence boundary for cleaner chunks
        if end < len(normalized):
            for punct in ".!?":
                boundary = normalized.rfind(punct + " ", start, end)
                if boundary != -1 and boundary > start + chunk_size // 2:
                    end = boundary + 1
                    break

        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start += step

    return chunks


def get_chroma_client() -> Any:
    """Get or create a persistent Chroma client."""
    _ensure_chroma()
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    return chromadb.PersistentClient(path=str(CHROMA_DIR))


def get_collection(project_id: int) -> Any:
    """Get or create a Chroma collection for a project."""
    _ensure_chroma()
    client = get_chroma_client()
    name = _collection_name_for_project(project_id)
    return client.get_or_create_collection(name=name)


def rebuild_collection(project_id: int) -> Any:
    """Delete and recreate a project's Chroma collection."""
    _ensure_chroma()
    client = get_chroma_client()
    name = _collection_name_for_project(project_id)
    try:
        client.delete_collection(name)
    except Exception:
        pass
    return client.get_or_create_collection(name=name)


def ingest_file(
    path: Path,
    project_id: int,
    character_id: Optional[int] = None,
    commit_id: Optional[int] = None,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> Dict[str, Any]:
    """Ingest a single file into a project-specific Chroma collection.

    Returns a summary dict with counts and IDs.
    """
    text = extract_text_from_file(path)
    chunks = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    if not chunks:
        return {
            "project_id": project_id,
            "source_file": path.name,
            "chunks_stored": 0,
            "collection_name": _collection_name_for_project(project_id),
        }

    collection = get_collection(project_id)
    documents: List[str] = []
    ids: List[str] = []
    metadatas: List[Dict[str, Any]] = []

    for index, chunk in enumerate(chunks, start=1):
        documents.append(chunk)
        doc_id = f"{path.stem}::chunk_{index}"
        ids.append(doc_id)
        meta: Dict[str, Any] = {
            "source_file": path.name,
            "chunk_index": index,
            "character_count": len(chunk),
            "project_id": project_id,
        }
        if character_id is not None:
            meta["character_id"] = character_id
        if commit_id is not None:
            meta["commit_id"] = commit_id
        metadatas.append(meta)

    collection.upsert(documents=documents, ids=ids, metadatas=metadatas)

    logger.info(
        "Ingested %s into project %s: %d chunks",
        path.name,
        project_id,
        len(documents),
    )

    return {
        "project_id": project_id,
        "source_file": path.name,
        "chunks_stored": len(documents),
        "collection_name": _collection_name_for_project(project_id),
    }


def ingest_directory(
    directory: Path,
    project_id: int,
    extensions: Optional[List[str]] = None,
    character_id: Optional[int] = None,
    commit_id: Optional[int] = None,
) -> Dict[str, Any]:
    """Ingest all supported files in a directory into a project collection."""
    from backend.config import SUPPORTED_UPLOAD_EXTENSIONS

    if extensions is None:
        extensions = list(SUPPORTED_UPLOAD_EXTENSIONS)

    collection = rebuild_collection(project_id)
    files_found = 0
    files_processed = 0
    total_chunks = 0

    for ext in extensions:
        for path in sorted(directory.glob(f"*{ext}")):
            files_found += 1
            try:
                result = ingest_file(
                    path=path,
                    project_id=project_id,
                    character_id=character_id,
                    commit_id=commit_id,
                )
                total_chunks += result["chunks_stored"]
                files_processed += 1
            except Exception as exc:
                logger.warning("Failed to ingest %s: %s", path.name, exc)

    return {
        "project_id": project_id,
        "directory": str(directory),
        "files_found": files_found,
        "files_processed": files_processed,
        "chunks_stored": total_chunks,
        "collection_name": _collection_name_for_project(project_id),
    }


def search_lore(
    query: str,
    project_id: int,
    n_results: int = 3,
    character_id: Optional[int] = None,
    commit_id: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """Search a project-specific Chroma collection for relevant lore chunks.

    Optional character_id / commit_id can be used for future metadata filtering.
    """
    if not query.strip():
        return []

    collection = get_collection(project_id)
    collection_count = collection.count()
    if collection_count == 0:
        return []

    result_count = min(max(1, n_results), collection_count)

    # Build optional filter
    where: Optional[Dict[str, Any]] = None
    if character_id is not None and commit_id is not None:
        where = {
            "$and": [
                {"project_id": project_id},
                {"character_id": character_id},
                {"commit_id": commit_id},
            ]
        }
    elif character_id is not None:
        where = {"$and": [{"project_id": project_id}, {"character_id": character_id}]}
    elif commit_id is not None:
        where = {"$and": [{"project_id": project_id}, {"commit_id": commit_id}]}

    kwargs: Dict[str, Any] = {
        "query_texts": [query],
        "n_results": result_count,
    }
    if where is not None:
        kwargs["where"] = where

    results = collection.query(**kwargs)

    documents = results.get("documents", [[]])[0]
    metadatas = results.get("metadatas", [[]])[0]
    distances = results.get("distances", [[]])[0]

    matches: List[Dict[str, Any]] = []
    for index, document in enumerate(documents):
        metadata = metadatas[index] if index < len(metadatas) else {}
        distance = distances[index] if index < len(distances) else None
        matches.append(
            {
                "text": document,
                "metadata": metadata or {},
                "distance": distance,
            }
        )

    return matches
